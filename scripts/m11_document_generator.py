#!/usr/bin/env python3
"""
M11 Protocol Document Generator

Generates an ICH M11 CeSHarP-compliant protocol document (Word/.docx)
from a USDM v4.0.0 JSON study definition.

Usage:
    python m11_document_generator.py \
        --input study_definition.json \
        --output protocol_m11.docx
"""

import json
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("ERROR: python-docx is required. Install with: pip install python-docx")
    raise

from usdm_utils import (
    get_version_and_design,
    sort_linked_list,
    get_study_title,
    get_study_id,
    get_sponsor_info,
    get_registry_id,
    get_criterion_text,
    get_enrollment_number,
)


def load_usdm(path: str) -> dict:
    """Load USDM JSON."""
    with open(path) as f:
        return json.load(f)


def add_title_page(doc: Document, version: dict, design: dict):
    """Generate the M11 Title Page."""
    sponsor = get_sponsor_info(version)
    phase = design.get("studyPhase", {}).get("standardCode", {}).get("decode", "TBD")
    title = get_study_title(version, "Official")
    version_id = version.get("versionIdentifier", "1.0")

    doc.add_heading("CLINICAL STUDY PROTOCOL", level=0)
    doc.add_paragraph("")

    # Protocol metadata table
    registry_id = get_registry_id(version)
    fields = [
        ("Protocol Title", title),
        ("Protocol Identifier", sponsor["protocol_number"]),
        ("Sponsor", sponsor["name"]),
        ("Study Phase", phase),
        ("Protocol Version", version_id),
        ("Registry Identifier", registry_id),
    ]

    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Table Grid"

    for i, (label, value) in enumerate(fields):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)
        cell_label.text = label
        cell_value.text = str(value) if value else "TBD"
        for paragraph in cell_label.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_page_break()


def add_synopsis(doc: Document, version: dict, design: dict):
    """Generate the Protocol Synopsis section."""
    doc.add_heading("Protocol Synopsis", level=1)

    sponsor = get_sponsor_info(version)
    title = get_study_title(version, "Official")
    phase = design.get("studyPhase", {}).get("standardCode", {}).get("decode", "TBD")

    # Model (v4.0.0: design.model, not interventionModel)
    model = design.get("model", design.get("interventionModel", {}))
    int_model = model.get("decode", "TBD") if isinstance(model, dict) else "TBD"

    # Arms
    arms = design.get("arms", design.get("studyArms", []))

    # Enrollment from population (singular)
    enrollment = get_enrollment_number(design)

    # Study type
    study_type = design.get("studyType", {}).get("decode", "TBD")

    # Blinding
    blinding = design.get("blindingSchema", {})
    blinding_text = ""
    if isinstance(blinding, dict):
        std_code = blinding.get("standardCode", {})
        blinding_text = std_code.get("decode", "TBD") if isinstance(std_code, dict) else "TBD"

    synopsis_items = [
        ("Sponsor", sponsor["name"]),
        ("Protocol Number", sponsor["protocol_number"]),
        ("Study Title", title),
        ("Study Phase", phase),
        ("Study Type", study_type),
        ("Intervention Model", int_model),
        ("Blinding", blinding_text),
        ("Number of Arms", str(len(arms))),
        ("Planned Enrollment", enrollment),
    ]

    table = doc.add_table(rows=len(synopsis_items), cols=2)
    table.style = "Table Grid"

    for i, (label, value) in enumerate(synopsis_items):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value if value else "TBD"
        for paragraph in table.cell(i, 0).paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph("")

    # Arms summary
    if arms:
        doc.add_heading("Treatment Arms", level=3)
        for arm in arms:
            arm_type = arm.get("type", arm.get("armType", {}))
            type_text = arm_type.get("decode", "") if isinstance(arm_type, dict) else ""
            arm_label = arm.get("label", arm.get("name", "TBD"))
            arm_desc = arm.get("description", "")
            doc.add_paragraph(
                f"{arm_label} ({type_text}): {arm_desc}",
                style="List Bullet"
            )

    # Indications
    indications = design.get("indications", [])
    if indications:
        doc.add_heading("Indication(s)", level=3)
        for ind in indications:
            doc.add_paragraph(
                f"{ind.get('label', ind.get('name', ''))}: {ind.get('description', '')}",
                style="List Bullet"
            )

    doc.add_page_break()


def add_trial_design(doc: Document, design: dict):
    """Generate Section 1: Trial Design."""
    doc.add_heading("1. Introduction", level=1)
    doc.add_heading("1.1 Trial Design", level=2)

    arms = design.get("arms", design.get("studyArms", []))
    epochs = sort_linked_list(design.get("epochs", design.get("studyEpochs", [])))

    if arms and epochs:
        doc.add_paragraph(
            "The following table summarizes the trial design as a matrix of arms and epochs:"
        )

        table = doc.add_table(rows=len(arms) + 1, cols=len(epochs) + 1)
        table.style = "Table Grid"

        # Header row
        table.cell(0, 0).text = "Arm \\ Epoch"
        for j, epoch in enumerate(epochs):
            table.cell(0, j + 1).text = epoch.get("label", epoch.get("name", ""))

        # Data rows
        cells = design.get("studyCells", [])
        elements = {e["id"]: e for e in design.get("elements", design.get("studyElements", []))}

        for i, arm in enumerate(arms):
            table.cell(i + 1, 0).text = arm.get("label", arm.get("name", ""))
            for j, epoch in enumerate(epochs):
                matching = [
                    c for c in cells
                    if c.get("armId") == arm["id"] and c.get("epochId") == epoch["id"]
                ]
                elem_names = []
                for cell in matching:
                    for eid in cell.get("elementIds", []):
                        elem = elements.get(eid, {})
                        elem_names.append(elem.get("label", elem.get("name", "")))
                table.cell(i + 1, j + 1).text = "\n".join(elem_names)

    doc.add_paragraph("")


def add_objectives_endpoints(doc: Document, design: dict):
    """Generate Section 2: Objectives and Endpoints.

    In v4.0.0, endpoints are embedded in each Objective, not a separate array.
    """
    doc.add_heading("2. Trial Objectives and Endpoints", level=1)

    objectives = design.get("objectives", [])

    for level_name in ["Primary", "Secondary", "Exploratory"]:
        level_objs = [
            o for o in objectives
            if level_name.lower() in o.get("level", {}).get("decode", "").lower()
        ]
        if not level_objs:
            continue

        doc.add_heading(f"2.x {level_name} Objective(s)", level=2)

        for obj in level_objs:
            text = obj.get("text", "")
            doc.add_paragraph(f"Objective: {text}")

            # Endpoints embedded in the objective (v4.0.0)
            endpoints = obj.get("endpoints", [])
            if endpoints:
                doc.add_paragraph("Associated Endpoint(s):")
                for ep in endpoints:
                    doc.add_paragraph(ep.get("text", ""), style="List Bullet")
            doc.add_paragraph("")


def add_eligibility(doc: Document, version: dict, design: dict):
    """Generate Section 4: Trial Population / Eligibility.

    In v4.0.0, criterion text is resolved via criterionItemId.
    """
    doc.add_heading("4. Trial Population", level=1)

    criteria = design.get("eligibilityCriteria", [])
    criterion_items = version.get("eligibilityCriterionItems", [])

    inclusion = [c for c in criteria if "inclusion" in c.get("category", {}).get("decode", "").lower()]
    exclusion = [c for c in criteria if "exclusion" in c.get("category", {}).get("decode", "").lower()]

    doc.add_heading("4.1 Inclusion Criteria", level=2)
    doc.add_paragraph("Subjects must meet ALL of the following criteria to be eligible:")
    for i, c in enumerate(inclusion, 1):
        text = get_criterion_text(c, criterion_items)
        doc.add_paragraph(f"{i}. {text}")

    doc.add_paragraph("")
    doc.add_heading("4.2 Exclusion Criteria", level=2)
    doc.add_paragraph("Subjects meeting ANY of the following criteria are excluded:")
    for i, c in enumerate(exclusion, 1):
        text = get_criterion_text(c, criterion_items)
        doc.add_paragraph(f"{i}. {text}")


def add_interventions(doc: Document, version: dict, design: dict):
    """Generate Section 5: Trial Interventions.

    In v4.0.0, studyInterventions are at the version level, referenced
    from design via studyInterventionIds[].
    """
    doc.add_heading("5. Trial Intervention(s)", level=1)

    # Get interventions from version level
    all_interventions = version.get("studyInterventions", [])
    intv_ids = design.get("studyInterventionIds", [])

    # Filter to those referenced by design, or use all if no IDs specified
    if intv_ids:
        interventions = [iv for iv in all_interventions if iv.get("id") in intv_ids]
    else:
        interventions = all_interventions

    if not interventions:
        doc.add_paragraph("No interventions defined.")
        return

    table = doc.add_table(rows=len(interventions) + 1, cols=3)
    table.style = "Table Grid"

    headers = ["Intervention", "Role", "Description"]
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
        for p in table.cell(0, j).paragraphs:
            for run in p.runs:
                run.bold = True

    for i, intv in enumerate(interventions, 1):
        table.cell(i, 0).text = intv.get("label", intv.get("name", ""))
        table.cell(i, 1).text = intv.get("role", {}).get("decode", "")
        table.cell(i, 2).text = intv.get("description", "")


def add_schedule_of_activities(doc: Document, design: dict):
    """Generate Section 6: Schedule of Activities."""
    doc.add_heading("6. Schedule of Activities", level=1)

    encounters = sort_linked_list(design.get("encounters", []))
    activities = sort_linked_list(design.get("activities", []))

    if not encounters or not activities:
        doc.add_paragraph("Schedule of Activities to be defined.")
        return

    doc.add_paragraph(
        "The following table summarizes the key assessments and their timing. "
        "Refer to individual sections for detailed procedures."
    )

    # SoA table: rows = activities, columns = encounters
    table = doc.add_table(rows=len(activities) + 1, cols=len(encounters) + 1)
    table.style = "Table Grid"

    # Header
    table.cell(0, 0).text = "Assessment"
    for j, enc in enumerate(encounters):
        cell = table.cell(0, j + 1)
        cell.text = enc.get("label", enc.get("name", ""))
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)

    # Activity rows
    for i, act in enumerate(activities, 1):
        table.cell(i, 0).text = act.get("label", act.get("name", ""))

    doc.add_paragraph("")
    doc.add_paragraph(
        "Note: Detailed timing and windows are defined in the full protocol.",
        style="Intense Quote"
    )


def generate_m11(input_path: str, output_path: str):
    """Main generation function."""
    data = load_usdm(input_path)
    version, design = get_version_and_design(data)

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)

    # Build document
    add_title_page(doc, version, design)
    add_synopsis(doc, version, design)
    add_trial_design(doc, design)
    add_objectives_endpoints(doc, design)
    add_eligibility(doc, version, design)
    add_interventions(doc, version, design)
    add_schedule_of_activities(doc, design)

    # Save
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"  ✓ M11 protocol document saved to: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description="Generate M11-compliant protocol document from USDM v4.0.0 JSON"
    )
    parser.add_argument("--input", "-i", required=True, help="Path to USDM JSON file")
    parser.add_argument("--output", "-o", required=True, help="Output .docx file path")
    args = parser.parse_args()

    generate_m11(args.input, args.output)


if __name__ == "__main__":
    main()
